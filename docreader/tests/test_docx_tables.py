import importlib.util
import io
import sys
import types
import unittest
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from unittest.mock import patch

from docx import Document as WordDocument


def _load_docx_parser():
    """Load docx_parser without triggering the heavy package __init__.

    ``docreader/parser/__init__.py`` imports doc_parser -> textract, a heavy
    dependency unrelated to DOCX parsing. Registering the packages as bare
    namespaces lets us import only the modules docx_parser actually needs.
    """
    root = Path(__file__).resolve().parents[2]
    docreader_pkg = types.ModuleType("docreader")
    docreader_pkg.__path__ = [str(root / "docreader")]
    sys.modules.setdefault("docreader", docreader_pkg)
    parser_pkg = types.ModuleType("docreader.parser")
    parser_pkg.__path__ = [str(root / "docreader" / "parser")]
    sys.modules["docreader.parser"] = parser_pkg

    spec = importlib.util.spec_from_file_location(
        "docreader.parser.docx_parser", root / "docreader" / "parser" / "docx_parser.py"
    )
    module = importlib.util.module_from_spec(spec)
    sys.modules["docreader.parser.docx_parser"] = module
    spec.loader.exec_module(module)
    return module


docx_parser = _load_docx_parser()
DocxParser = docx_parser.DocxParser


def _parse(content):
    """Parse a DOCX through the real Docx processor.

    The production path uses a ProcessPoolExecutor backed by a multiprocessing
    Manager. Some CI sandboxes forbid POSIX semaphores, so run the same page
    task pool on threads and swap the manager for a plain list - parse
    behavior is identical, only the execution backend changes.
    """
    class _FakeManager:
        def __enter__(self):
            self._items = []
            return self

        def __exit__(self, *exc):
            return False

        def list(self):
            return self._items

    with patch.object(docx_parser, "Manager", _FakeManager), patch.object(
        docx_parser, "ProcessPoolExecutor", ThreadPoolExecutor
    ):
        return DocxParser(max_pages=100).parse_into_text(content)


def _make_docx(add_table):
    doc = WordDocument()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class DocxTableContentTest(unittest.TestCase):
    """Regression test: DOCX tables must be kept in the parsed text.

    Docx.__call__ returns tables separately from the text lines, and
    parse_into_text used to drop them, silently losing all table content.
    """

    def _docx_with_table(self, cells):
        doc = WordDocument()
        doc.add_paragraph("Introduction paragraph")
        table = doc.add_table(rows=len(cells), cols=len(cells[0]))
        for r, row in enumerate(cells):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_table_content_is_kept_in_parsed_text(self):
        content = self._docx_with_table(
            [["City", "Population"], ["Beijing", "21.5M"]]
        )
        document = _parse(content)
        self.assertIn("Introduction paragraph", document.content)
        for cell in ("City", "Population", "Beijing", "21.5M"):
            self.assertIn(cell, document.content)
        # Rendered as a GFM markdown table: header + delimiter row.
        self.assertIn("| City | Population |", document.content)
        self.assertIn("| --- | --- |", document.content)

    def test_table_only_document_is_not_empty(self):
        doc = WordDocument()
        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "Header"
        table.cell(1, 0).text = "Value"
        buf = io.BytesIO()
        doc.save(buf)

        document = _parse(buf.getvalue())
        self.assertIn("Header", document.content)
        self.assertIn("Value", document.content)

    def test_pipe_in_cell_does_not_break_table(self):
        doc = WordDocument()
        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "A|B"
        table.cell(1, 0).text = "C"
        buf = io.BytesIO()
        doc.save(buf)

        document = _parse(buf.getvalue())
        self.assertIn(r"A\|B", document.content)

    def test_tables_to_gfm_markdown_helper(self):
        # Unit-level coverage for the renderer, using the exact tuple shape
        # Docx._process_tables returns: ((None, table_html), alt).
        html = (
            "<table><tr><td>A</td><td>B</td></tr>"
            "<tr><td>C</td><td>D</td></tr></table>"
        )
        rendered = DocxParser(max_pages=100)._tables_to_gfm_markdown(
            [((None, html), "")]
        )
        self.assertEqual(
            rendered,
            "| A | B |\n| --- | --- |\n| C | D |",
        )

    def test_tables_to_gfm_markdown_skips_empty(self):
        rendered = DocxParser(max_pages=100)._tables_to_gfm_markdown(
            [((None, ""), ""), ((None, None), "")]
        )
        self.assertEqual(rendered, "")


if __name__ == "__main__":
    unittest.main()