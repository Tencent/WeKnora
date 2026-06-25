import io
import unittest

from docx import Document

from docreader.parser.docx_table_parser import (
    TableRenderState,
    render_structured_table,
    structured_docx_tables_to_markdown,
    table_to_matrix,
)
from docreader.parser.markitdown_parser import StdMarkitdownParser


def _docx_bytes(doc: Document) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class DocxTableParserTest(unittest.TestCase):
    def test_matrix_recovers_horizontal_and_vertical_merges(self):
        doc = Document()
        table = doc.add_table(rows=4, cols=3)
        table.cell(0, 0).text = "Report"
        table.cell(0, 0).merge(table.cell(0, 2))
        table.cell(1, 0).text = "Category"
        table.cell(1, 1).text = "Issue"
        table.cell(1, 2).text = "Resolution"
        table.cell(2, 0).text = "Power"
        table.cell(2, 0).merge(table.cell(3, 0))
        table.cell(2, 1).text = "Will not start"
        table.cell(2, 2).text = "Replace battery"
        table.cell(3, 1).text = "Shuts down"
        table.cell(3, 2).text = "Clean contacts"

        parsed = Document(io.BytesIO(_docx_bytes(doc)))
        matrix = table_to_matrix(parsed.tables[0])

        self.assertEqual(matrix[0][0].text, "Report")
        self.assertIs(matrix[0][0], matrix[0][2])
        self.assertEqual(matrix[0][0].colspan, 3)
        self.assertEqual(matrix[2][0].text, "Power")
        self.assertIs(matrix[2][0], matrix[3][0])
        self.assertEqual(matrix[2][0].rowspan, 2)

    def test_render_structured_rows_with_context_and_inherited_category(self):
        doc = Document()
        table = doc.add_table(rows=4, cols=3)
        table.cell(0, 0).text = "Product A"
        table.cell(0, 0).merge(table.cell(0, 2))
        table.cell(1, 0).text = "Category"
        table.cell(1, 1).text = "Symptom"
        table.cell(1, 2).text = "Action"
        table.cell(2, 0).text = "Power"
        table.cell(2, 1).text = "Will not start"
        table.cell(2, 2).text = "Replace battery"
        table.cell(3, 1).text = "Shuts down"
        table.cell(3, 2).text = "Clean contacts"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        self.assertIn("## Structured tables", rendered)
        self.assertIn("Table 1 context: Product A", rendered)
        self.assertIn("Category: Power", rendered)
        self.assertIn("Symptom: Will not start", rendered)
        self.assertIn("Action: Replace battery", rendered)
        self.assertIn("Symptom: Shuts down", rendered)
        self.assertIn("Action: Clean contacts", rendered)

    def test_full_width_note_after_data_does_not_become_context(self):
        doc = Document()
        table = doc.add_table(rows=5, cols=2)
        table.cell(0, 0).text = "Topic"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 0).text = "Question"
        table.cell(1, 1).text = "Answer"
        table.cell(2, 0).text = "Q1"
        table.cell(2, 1).text = "A1"
        table.cell(3, 0).text = "Note: applies to the rows above"
        table.cell(3, 0).merge(table.cell(3, 1))
        table.cell(4, 0).text = "Q2"
        table.cell(4, 1).text = "A2"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        self.assertIn("Table 1 note: Note: applies to the rows above", rendered)
        self.assertIn("Question: Q2", rendered)
        q2_block = rendered[rendered.index("Table 1 row 5") :].split("\n\n")[0]
        self.assertIn("Context: Topic", q2_block)
        self.assertNotIn("Note: applies to the rows above", q2_block)

    def test_long_first_data_row_is_not_treated_as_header(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Operations"
        table.cell(0, 1).text = "Daily cleanup"
        table.cell(0, 2).text = "Run the cleanup job, verify the report, and escalate if the row count is unexpected."
        table.cell(1, 1).text = "Weekly cleanup"
        table.cell(1, 2).text = "Archive old records"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        self.assertIn("Column 1: Operations", rendered)
        self.assertIn("Column 2: Daily cleanup", rendered)
        row2_block = rendered[rendered.index("Table 1 row 2") :].split("\n\n")[0]
        self.assertIn("Column 1: Operations", row2_block)
        self.assertIn("Column 2: Weekly cleanup", row2_block)
        self.assertNotIn("Operations, Daily cleanup", rendered)

    def test_numbered_first_data_row_is_not_treated_as_header(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "2. Category"
        table.cell(0, 1).text = "2.1 Item"
        table.cell(0, 2).text = "Action"
        table.cell(1, 1).text = "2.2 Other item"
        table.cell(1, 2).text = "Other action"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        self.assertIn("Column 1: 2. Category", rendered)
        self.assertIn("Column 2: 2.1 Item", rendered)
        row2_block = rendered[rendered.index("Table 1 row 2") :].split("\n\n")[0]
        self.assertIn("Column 1: 2. Category", row2_block)
        self.assertIn("Column 2: 2.2 Other item", row2_block)
        self.assertNotIn("Table 1 columns: 2. Category", rendered)

    def test_continuation_table_inherits_headers_and_context(self):
        doc = Document()
        first = doc.add_table(rows=3, cols=3)
        first.cell(0, 0).text = "Product A"
        first.cell(0, 0).merge(first.cell(0, 2))
        first.cell(1, 0).text = "Category"
        first.cell(1, 1).text = "Issue"
        first.cell(1, 2).text = "Action"
        first.cell(2, 0).text = "Power"
        first.cell(2, 1).text = "Will not start"
        first.cell(2, 2).text = "Replace battery"

        second = doc.add_table(rows=2, cols=3)
        second.cell(0, 0).text = "Power"
        second.cell(0, 1).text = "Shuts down"
        second.cell(0, 2).text = "Clean contacts"
        second.cell(1, 1).text = "No sound"
        second.cell(1, 2).text = "Check volume"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        self.assertIn("### Table 2", rendered)
        table2 = rendered[rendered.index("### Table 2") :]
        self.assertIn("Context: Product A", table2)
        self.assertIn("Category: Power", table2)
        self.assertIn("Issue: Shuts down", table2)
        self.assertIn("Action: Clean contacts", table2)
        self.assertIn("Issue: No sound", table2)
        self.assertIn("Action: Check volume", table2)
        self.assertNotIn("Column 1: Power", table2)

    def test_continuation_table_inherits_previous_row_values(self):
        doc = Document()
        first = doc.add_table(rows=3, cols=3)
        first.cell(0, 0).text = "Category"
        first.cell(0, 1).text = "Issue"
        first.cell(0, 2).text = "Action"
        first.cell(1, 0).text = "Power"
        first.cell(1, 1).text = "Will not start"
        first.cell(1, 2).text = "Replace battery"
        first.cell(2, 0).text = "Power"
        first.cell(2, 1).text = "Shuts down"
        first.cell(2, 2).text = "Clean contacts"

        second = doc.add_table(rows=1, cols=3)
        second.cell(0, 1).text = "Overheats"
        second.cell(0, 2).text = "Let it cool"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        table2 = rendered[rendered.index("### Table 2") :]
        self.assertIn("Category: Power", table2)
        self.assertIn("Issue: Overheats", table2)
        self.assertIn("Action: Let it cool", table2)

    def test_short_continuation_row_reusing_prior_value_is_not_header(self):
        doc = Document()
        first = doc.add_table(rows=2, cols=3)
        first.cell(0, 0).text = "Category"
        first.cell(0, 1).text = "Issue"
        first.cell(0, 2).text = "Action"
        first.cell(1, 0).text = "Power"
        first.cell(1, 1).text = "Start"
        first.cell(1, 2).text = "Replace"

        second = doc.add_table(rows=1, cols=3)
        second.cell(0, 0).text = "Power"
        second.cell(0, 1).text = "Stop"
        second.cell(0, 2).text = "Clean"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        table2 = rendered[rendered.index("### Table 2") :]
        self.assertIn("Category: Power", table2)
        self.assertIn("Issue: Stop", table2)
        self.assertIn("Action: Clean", table2)
        self.assertNotIn("Table 2 columns: Power", table2)

    def test_continuation_state_resets_on_new_context_row(self):
        doc = Document()
        first = doc.add_table(rows=2, cols=2)
        first.cell(0, 0).text = "Name"
        first.cell(0, 1).text = "Value"
        first.cell(1, 0).text = "Timeout"
        first.cell(1, 1).text = "30 seconds"

        second = doc.add_table(rows=2, cols=2)
        second.cell(0, 0).text = "New section"
        second.cell(0, 0).merge(second.cell(0, 1))
        second.cell(1, 0).text = "Retries"
        second.cell(1, 1).text = "3"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        table2 = rendered[rendered.index("### Table 2") :]
        self.assertIn("Table 2 context: New section", table2)
        self.assertIn("Column 1: Retries", table2)
        self.assertIn("Column 2: 3", table2)
        self.assertNotIn("Name: Retries", table2)

    def test_same_width_table_with_own_header_does_not_inherit(self):
        doc = Document()
        first = doc.add_table(rows=2, cols=2)
        first.cell(0, 0).text = "Name"
        first.cell(0, 1).text = "Value"
        first.cell(1, 0).text = "Timeout"
        first.cell(1, 1).text = "30 seconds"

        second = doc.add_table(rows=2, cols=2)
        second.cell(0, 0).text = "Metric"
        second.cell(0, 1).text = "Score"
        second.cell(1, 0).text = "Quality"
        second.cell(1, 1).text = "95"

        rendered = structured_docx_tables_to_markdown(_docx_bytes(doc))

        table2 = rendered[rendered.index("### Table 2") :]
        self.assertIn("Table 2 columns: Metric, Score", table2)
        self.assertIn("Metric: Quality", table2)
        self.assertIn("Score: 95", table2)
        self.assertNotIn("Name: Quality", table2)

    def test_render_structured_table_can_be_called_without_prior_state(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Timeout"
        table.cell(1, 1).text = "30 seconds"

        rendered, state = render_structured_table(table, 1)

        self.assertIsInstance(state, TableRenderState)
        self.assertIn("Name: Timeout", rendered)

    def test_markitdown_docx_parser_appends_structured_table_section(self):
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Timeout"
        table.cell(1, 1).text = "30 seconds"

        parsed = StdMarkitdownParser(file_type="docx").parse_into_text(
            _docx_bytes(doc)
        )

        self.assertIn("Before table", parsed.content)
        self.assertIn("## Structured tables", parsed.content)
        self.assertIn("Name: Timeout", parsed.content)
        self.assertIn("Value: 30 seconds", parsed.content)


if __name__ == "__main__":
    unittest.main()
