import io
import unittest
import zipfile
from pathlib import Path

from docx import Document as WordDocument
from lxml import etree

from docreader.parser.docx2_parser import Docx2Parser
from docreader.parser.docx_merge import fill_vertical_merged_cells_docx
from docreader.parser.markitdown_parser import StdMarkitdownParser

FIXTURE = Path(__file__).with_name("fixtures") / "issue_2634_vertical_merge.docx"
METHOD_PREFIX = "检测方法：提取外周血基因组 DNA"
PROJECT_CODES = ("Q0101", "Q0102", "Q0103", "Q0104")
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class DocxVerticalMergeFillTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.fixture_bytes = FIXTURE.read_bytes()

    def test_fixture_is_a_real_docx_with_vertical_merge(self):
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(self.fixture_bytes)))
        with zipfile.ZipFile(io.BytesIO(self.fixture_bytes)) as archive:
            self.assertIsNone(archive.testzip())
            root = etree.fromstring(archive.read("word/document.xml"))

        vertical_merges = root.xpath(
            ".//w:tbl[1]//w:vMerge", namespaces={"w": WORD_NAMESPACE}
        )
        merge_values = [
            node.get(f"{{{WORD_NAMESPACE}}}val", "continue") for node in vertical_merges
        ]
        self.assertEqual(
            merge_values,
            ["restart", "continue", "continue", "continue"],
        )

    def test_fill_repeats_master_content_in_every_covered_row(self):
        filled = fill_vertical_merged_cells_docx(self.fixture_bytes)
        document = WordDocument(io.BytesIO(filled))
        table = document.tables[0]

        for row_index, project_code in enumerate(PROJECT_CODES, start=1):
            self.assertEqual(table.cell(row_index, 0).text, project_code)
            self.assertIn(METHOD_PREFIX, table.cell(row_index, 4).text)

        with zipfile.ZipFile(io.BytesIO(filled)) as archive:
            root = etree.fromstring(archive.read("word/document.xml"))
        self.assertEqual(
            root.xpath("count(.//w:vMerge)", namespaces={"w": WORD_NAMESPACE}),
            0.0,
        )

    def test_non_merged_equal_cells_are_not_collapsed(self):
        filled = fill_vertical_merged_cells_docx(self.fixture_bytes)
        document = WordDocument(io.BytesIO(filled))
        control_row = document.tables[1].rows[1]
        self.assertEqual(len(control_row.cells), 3)
        self.assertEqual(
            [cell.text for cell in control_row.cells],
            ["相同值", "相同值", "独立值"],
        )

    def test_non_docx_bytes_pass_through_unchanged(self):
        content = b"not a docx package"
        self.assertIs(fill_vertical_merged_cells_docx(content), content)

    def test_docx_without_vertical_merges_passes_through_unchanged(self):
        document = WordDocument()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "left"
        table.cell(0, 1).text = "right"
        output = io.BytesIO()
        document.save(output)
        content = output.getvalue()

        self.assertIs(fill_vertical_merged_cells_docx(content), content)


class Issue2634ParserRegressionTest(unittest.TestCase):
    def test_markitdown_output_repeats_merged_value_for_every_data_row(self):
        document = StdMarkitdownParser(
            file_name=FIXTURE.name, file_type="docx"
        ).parse_into_text(FIXTURE.read_bytes())

        data_lines = {
            project_code: next(
                line for line in document.content.splitlines() if project_code in line
            )
            for project_code in PROJECT_CODES
        }
        for line in data_lines.values():
            self.assertIn(METHOD_PREFIX, line)
        self.assertEqual(document.content.count(METHOD_PREFIX), len(PROJECT_CODES))
        self.assertIn("| 相同值 | 相同值 | 独立值 |", document.content)

    def test_full_docx_entry_keeps_context_on_all_four_rows(self):
        document = Docx2Parser(
            file_name=FIXTURE.name, file_type="docx"
        ).parse_into_text(FIXTURE.read_bytes())

        for project_code in PROJECT_CODES:
            data_line = next(
                line for line in document.content.splitlines() if project_code in line
            )
            self.assertIn(METHOD_PREFIX, data_line)
        self.assertEqual(document.content.count(METHOD_PREFIX), len(PROJECT_CODES))


if __name__ == "__main__":
    unittest.main()
